# -*- coding: utf-8 -*-
import copy
import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

TEMPLATE = r"C:\Users\a.abboud\Downloads\Weekly Planner Template x4 - Final.docx"
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(8)

def set_run_style(run):
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), FONT_NAME)
    rfonts.set(qn('w:hAnsi'), FONT_NAME)
    rfonts.set(qn('w:cs'), FONT_NAME)

def fill_cell_lines(cell, lines):
    """Fill an empty cell with one or more lines of plain text, matching template's compact style."""
    if not lines:
        return
    first_para = cell.paragraphs[0]
    for i, line in enumerate(lines):
        p = first_para if i == 0 else cell.add_paragraph()
        r = p.add_run(line)
        set_run_style(r)

def append_to_first_run_paragraph(cell, para_index, text):
    """Append text right after the existing single run in a paragraph (used for Beginning/On/Above rows).
    Some template columns are missing the trailing space after the label (e.g. 'Above:' vs 'Above: ') -
    pad our own inserted text instead of touching the template's existing label run."""
    p = cell.paragraphs[para_index]
    label_run_text = p.runs[0].text if p.runs else ''
    if label_run_text and not label_run_text.endswith(' '):
        text = ' ' + text
    r = p.add_run(text)
    set_run_style(r)

def check_boxes(cell, labels_to_check):
    """Flip unicode checkbox glyph for runs whose following-run text matches a target label."""
    for p in cell.paragraphs:
        runs = p.runs
        for i, r in enumerate(runs):
            if r.text == '☐' and i + 1 < len(runs):
                label = runs[i + 1].text.strip()
                for target in labels_to_check:
                    if target.lower() in label.lower():
                        r.text = '☑'

def fill_grade(doc, grade_text):
    for p in doc.paragraphs:
        if p.text.strip().startswith('Grade:'):
            r = p.add_run(grade_text)
            set_run_style(r)
            return

def fill_week_header(table, week_line, lesson_titles):
    row0 = table.rows[0]
    r = row0.cells[0].paragraphs[0].add_run(' ' + week_line)
    set_run_style(r)
    for i, title in enumerate(lesson_titles):
        r = row0.cells[1 + i].paragraphs[0].add_run(' ' + title)
        set_run_style(r)

ROW = {
    'standard': 1, 'literacy': 2, 'assessment': 3, 'walt': 4, 'wilf': 5,
    'vocab': 6, 'resources': 7, 'national_identity': 8, 'ai_digital': 9,
    'i_do': 10, 'we_do': 11, 'you_do': 12, 'closure': 14,
}
DIFF_ROW = 13
INCLUSION_ROW = 15
GT_ROW = 16

def fill_week(week_data, out_path):
    doc = docx.Document(TEMPLATE)
    table = doc.tables[0]

    fill_grade(doc, ' 9')
    fill_week_header(table, week_data['week_line'], [l['title'] for l in week_data['lessons']])

    for li, lesson in enumerate(week_data['lessons']):
        col = 1 + li
        for key, row_idx in ROW.items():
            cell = table.rows[row_idx].cells[col]
            val = lesson.get(key)
            if val is None:
                continue
            lines = val if isinstance(val, list) else [val]
            fill_cell_lines(cell, lines)

        diff_cell = table.rows[DIFF_ROW].cells[col]
        append_to_first_run_paragraph(diff_cell, 0, lesson['diff_beginning'])
        append_to_first_run_paragraph(diff_cell, 1, lesson['diff_on'])
        append_to_first_run_paragraph(diff_cell, 2, lesson['diff_above'])

        check_boxes(table.rows[INCLUSION_ROW].cells[col], lesson.get('inclusion', []))
        check_boxes(table.rows[GT_ROW].cells[col], lesson.get('gt', []))

    doc.save(out_path)
    print('Saved', out_path)
